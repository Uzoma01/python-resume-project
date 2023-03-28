from docx import Document

document = Document()

# Name and contact  details

name = input("What is your name? ")
email = input("What is your E-mail? ")

document.add_paragraph(name)
document.add_paragraph(email) 

# About me

document.add_heading("About Me")
about_me = input("Tell us about yourself. ")
document.add_paragraph(about_me)

# Work Experience

document.add_heading("Work Experience")
p = document.add_paragraph()

company = input("Enter Company Name: ")
from_date = input("From date: ")
to_date = input("To date: ")

p.add_run(company + "\n").bold = True
p.add_run(from_date + " - " + to_date + "\n").italic = True

experience_details = input("Describe your Experience: ")
p.add_run(experience_details)

# More Experiences

while True:
    has_more_experiences = input("Do you have more Experiences? Yes or No? ")
    if has_more_experiences.lower() == "yes":
        p = document.add_paragraph()       
        company = input("Enter Company Name: ")
        from_date = input("From date: ")
        to_date = input("To date: ")

        p.add_run(company + "\n").bold = True
        p.add_run(from_date + " - " + to_date + "\n").italic = True

        experience_details = input("Describe your Experience: ")
        p.add_run(experience_details)
    else:
        break

# Skills

document.add_heading("Skills")
skills = input("Enter a Skill: ")
document.add_paragraph(skills, style="List Bullet")

# More Skills

while True:
    has_more_skills = input("Do you have more skills? Yes or No? ")
    if has_more_skills.lower() == "yes":
        skills = input("Enter a Skill: ")
        document.add_paragraph(skills, style="List Bullet")
    else:
        break
    

document.save("resume.docx")


